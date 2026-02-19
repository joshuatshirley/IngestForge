"""Document formatters for outline export.

Generates formatted documents from mapped outlines with evidence.
Supports Markdown and DOCX output formats."""

from __future__ import annotations

from abc import ABC, abstractmethod
from datetime import datetime
from pathlib import Path
from typing import List

from ingestforge.core.export.outline_mapper import (
    MappedOutline,
    OutlinePoint,
    EvidenceMatch,
)
from ingestforge.core.logging import get_logger

logger = get_logger(__name__)
MAX_EVIDENCE_PER_POINT = 10
MAX_CONTENT_LENGTH = 500


class DocumentFormatter(ABC):
    """Abstract base for document formatters."""

    @abstractmethod
    def format(self, outline: MappedOutline) -> str:
        """Format outline to document content.

        Args:
            outline: Mapped outline with evidence

        Returns:
            Formatted document content
        """
        pass

    @abstractmethod
    def save(self, outline: MappedOutline, output_path: Path) -> None:
        """Save formatted document to file.

        Args:
            outline: Mapped outline with evidence
            output_path: Output file path
        """
        pass


class MarkdownFormatter(DocumentFormatter):
    """Format outlines to Markdown documents."""

    def __init__(
        self,
        include_citations: bool = True,
        include_metadata: bool = True,
    ) -> None:
        """Initialize formatter.

        Args:
            include_citations: Include citations section
            include_metadata: Include document metadata header
        """
        self.include_citations = include_citations
        self.include_metadata = include_metadata

    def format(self, outline: MappedOutline) -> str:
        """Format outline to Markdown.

        Args:
            outline: Mapped outline with evidence

        Returns:
            Markdown content string
        """
        parts: List[str] = []

        # Document header
        parts.append(self._format_header(outline))

        # Outline points with evidence
        for point in outline.points:
            parts.append(self._format_point(point, outline))

        # Citations section
        if self.include_citations:
            parts.append(self._format_citations(outline))

        return "\n".join(parts)

    def save(self, outline: MappedOutline, output_path: Path) -> None:
        """Save to Markdown file.

        Args:
            outline: Mapped outline
            output_path: Output .md file path
        """
        content = self.format(outline)
        output_path.write_text(content, encoding="utf-8")
        logger.info(f"Saved Markdown to: {output_path}")

    def _format_header(self, outline: MappedOutline) -> str:
        """Format document header.

        Args:
            outline: Mapped outline

        Returns:
            Header string
        """
        parts = [f"# {outline.title}\n"]

        if self.include_metadata:
            timestamp = datetime.now().strftime("%Y-%m-%d %H:%M")
            parts.append(f"\n*Generated: {timestamp}*\n")
            parts.append(f"*Points: {len(outline.points)}*\n")
            parts.append(f"*Evidence: {outline.get_total_evidence_count()}*\n")

        parts.append("\n---\n")
        return "".join(parts)

    def _format_point(self, point: OutlinePoint, outline: MappedOutline) -> str:
        """Format a single outline point with evidence.

        Args:
            point: Outline point
            outline: Full outline for evidence lookup

        Returns:
            Formatted point string
        """
        parts: List[str] = []

        # Point header
        prefix = "#" * (point.level + 1)  # +1 because title is #
        parts.append(f"\n{prefix} {point.title}\n")

        if point.description:
            parts.append(f"\n{point.description}\n")

        # Evidence for this point
        evidence = outline.get_evidence_for_point(point.id)[:MAX_EVIDENCE_PER_POINT]

        if evidence:
            parts.append("\n**Evidence:**\n")
            for match in evidence:
                parts.append(self._format_evidence(match))

        return "".join(parts)

    def _format_evidence(self, match: EvidenceMatch) -> str:
        """Format a single evidence match.

        Args:
            match: Evidence match

        Returns:
            Formatted evidence string
        """
        # Truncate long content
        content = match.chunk_content[:MAX_CONTENT_LENGTH]
        if len(match.chunk_content) > MAX_CONTENT_LENGTH:
            content += "..."

        citation = match.to_citation()
        citation_ref = f" [{citation}]" if citation else ""

        return f"\n> {content}{citation_ref}\n"

    def _format_citations(self, outline: MappedOutline) -> str:
        """Format citations section.

        Args:
            outline: Mapped outline

        Returns:
            Citations section string
        """
        sources: dict[str, int] = {}

        for matches in outline.evidence.values():
            for match in matches:
                source = match.source_file or "Unknown source"
                sources[source] = sources.get(source, 0) + 1

        if not sources:
            return ""

        parts = ["\n---\n\n## Sources\n"]

        for idx, (source, count) in enumerate(sorted(sources.items()), 1):
            ref_word = "reference" if count == 1 else "references"
            parts.append(f"\n{idx}. **{source}** - {count} {ref_word}")

        parts.append("\n")
        return "".join(parts)


class DocxFormatter(DocumentFormatter):
    """Format outlines to DOCX documents.

    Requires python-docx package.
    """

    def __init__(
        self,
        include_citations: bool = True,
        include_toc: bool = True,
    ) -> None:
        """Initialize formatter.

        Args:
            include_citations: Include citations section
            include_toc: Include table of contents
        """
        self.include_citations = include_citations
        self.include_toc = include_toc
        self._check_docx_available()

    def _check_docx_available(self) -> None:
        """Check if python-docx is available."""
        try:
            import docx  # noqa: F401
        except ImportError:
            logger.warning(
                "python-docx not installed. Install with: pip install python-docx"
            )

    def format(self, outline: MappedOutline) -> str:
        """Format outline to DOCX XML.

        Note: Returns empty string; use save() for DOCX output.

        Args:
            outline: Mapped outline

        Returns:
            Empty string (DOCX is binary format)
        """
        return ""

    def save(self, outline: MappedOutline, output_path: Path) -> None:
        """Save to DOCX file.

        Args:
            outline: Mapped outline
            output_path: Output .docx file path

        Raises:
            ImportError: If python-docx not installed
        """
        try:
            from docx import Document
            from docx.shared import Pt, Inches
            from docx.enum.style import WD_STYLE_TYPE
        except ImportError:
            raise ImportError(
                "python-docx required for DOCX export. "
                "Install with: pip install python-docx"
            )

        doc = Document()

        # Title
        doc.add_heading(outline.title, 0)

        # Metadata
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M")
        doc.add_paragraph(f"Generated: {timestamp}")
        doc.add_paragraph(f"Points: {len(outline.points)}")
        doc.add_paragraph(f"Evidence: {outline.get_total_evidence_count()}")

        doc.add_page_break()

        # Outline points
        for point in outline.points:
            self._add_point_to_doc(doc, point, outline)

        # Citations
        if self.include_citations:
            self._add_citations_to_doc(doc, outline)

        doc.save(str(output_path))
        logger.info(f"Saved DOCX to: {output_path}")

    def _add_point_to_doc(
        self, doc: "Document", point: OutlinePoint, outline: MappedOutline
    ) -> None:
        """Add outline point to document.

        Args:
            doc: Word document
            point: Outline point
            outline: Full outline for evidence
        """

        # Point heading
        doc.add_heading(point.title, point.level)

        if point.description:
            doc.add_paragraph(point.description)

        # Evidence
        evidence = outline.get_evidence_for_point(point.id)[:MAX_EVIDENCE_PER_POINT]

        if evidence:
            doc.add_paragraph("Evidence:", style="Intense Quote")

            for match in evidence:
                self._add_evidence_to_doc(doc, match)

    def _add_evidence_to_doc(self, doc: "Document", match: EvidenceMatch) -> None:
        """Add evidence to document.

        Args:
            doc: Word document
            match: Evidence match
        """
        content = match.chunk_content[:MAX_CONTENT_LENGTH]
        if len(match.chunk_content) > MAX_CONTENT_LENGTH:
            content += "..."

        para = doc.add_paragraph(content, style="Quote")

        citation = match.to_citation()
        if citation:
            para.add_run(f" [{citation}]").italic = True

    def _add_citations_to_doc(self, doc: "Document", outline: MappedOutline) -> None:
        """Add citations section to document.

        Args:
            doc: Word document
            outline: Mapped outline
        """
        sources: dict[str, int] = {}

        for matches in outline.evidence.values():
            for match in matches:
                source = match.source_file or "Unknown source"
                sources[source] = sources.get(source, 0) + 1

        if not sources:
            return

        doc.add_page_break()
        doc.add_heading("Sources", 1)

        for idx, (source, count) in enumerate(sorted(sources.items()), 1):
            ref_word = "reference" if count == 1 else "references"
            doc.add_paragraph(f"{idx}. {source} - {count} {ref_word}")


def get_formatter(
    output_format: str,
    include_citations: bool = True,
) -> DocumentFormatter:
    """Get formatter for output format.

    Args:
        output_format: Format type ("markdown", "md", "docx", "word")
        include_citations: Include citations section

    Returns:
        Appropriate DocumentFormatter instance

    Raises:
        ValueError: If format is unknown
    """
    format_lower = output_format.lower()

    if format_lower in ("markdown", "md"):
        return MarkdownFormatter(include_citations=include_citations)

    if format_lower in ("docx", "word"):
        return DocxFormatter(include_citations=include_citations)

    raise ValueError(f"Unknown output format: {output_format}")
